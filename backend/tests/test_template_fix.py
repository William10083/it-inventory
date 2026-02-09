"""
Test actualizado para verificar inserción de tabla con el nuevo template ID 8
"""
import sys
import os
sys.path.append(r'C:\Users\wvilca\Downloads\it_inventory 2\it_inventory\backend')

from database import SessionLocal
from models import DocumentTemplate
from pdf_generator import generate_batch_acta
from docx import Document

db = SessionLocal()

try:
    # Obtener template ID 8
    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == 8).first()
    
    if not template:
        print("ERROR: Template ID 8 not found")
        sys.exit(1)
    
    print("="*80)
    print("TESTING TABLE INSERTION")
    print("="*80)
    print(f"Template: {template.name}")
    print(f"Template Path: {template.file_path}")
    print(f"Template Type: {template.template_type}")
    
    # Datos de prueba
    employee_data = {
        'name': 'ADRIANA CASTRO',
        'dni': '60747728',
        'company': 'TRANSTOTAL'
    }
    
    devices_info = [
        {
            'type': 'LAPTOP',
            'brand': 'HP',
            'model': 'PROBOOK 440 G11',
            'serial': '5CD5401M1L',
            'hostname': 'TAM-LP001'
        },
        {
            'type': 'HEADPHONES',
            'brand': 'HP',
            'model': 'G2',
            'serial': '7CH4290H62',
            'hostname': '-'
        },
        {
            'type': 'KEYBOARD_MOUSE_KIT',
            'brand': 'HP',
            'model': 'HSA-A005K / HSA-A011M',
            'serial': '7CH43452CM',
            'hostname': '-'
        },
        {
            'type': 'BACKPACK',
            'brand': 'HP',
            'model': 'TRAVEL 18L',
            'serial': 'BACKPACK-36',
            'hostname': '-'
        },
        {
            'type': 'MONITOR',
            'brand': 'LENOVO',
            'model': 'THINKVISION E24-30',
            'serial': 'VNABFZ0H',
            'hostname': 'TAM-MN001'
        },
        {
            'type': 'CHARGER',
            'brand': 'HP',
            'model': 'TPN-DA15',
            'serial': 'CHARGER-HP-225',
            'hostname': '-'
        }
    ]
    
    print("\n" + "="*80)
    print("GENERATING TEST ACTA")
    print("="*80)
    
    test_acta_path = generate_batch_acta(
        assignment_id=9999,
        employee_name=employee_data['name'],
        devices_info=devices_info,
        employee_dni=employee_data['dni'],
        employee_company=employee_data['company'],
        template_path=template.file_path,
        template=template
    )
    
    print(f"\n[OK] Test acta generated at: {test_acta_path}")
    print(f"[OK] File exists: {os.path.exists(test_acta_path)}")
    
    # Verificar el contenido
    doc = Document(test_acta_path)
    
    print("\n" + "="*80)
    print("VERIFYING DOCUMENT CONTENT")
    print("="*80)
    
    # Buscar título
    for para in doc.paragraphs[:10]:
        if para.text.strip():
            print(f"Para: {para.text[:100]}")
    
    # Verificar si hay tablas
    print(f"\nNumber of tables in document: {len(doc.tables)}")
    
    if len(doc.tables) > 0:
        print("\n[OK] Tables found in document")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i}: {len(table.rows)} rows, {len(table.rows[0].cells) if table.rows else 0} cols")
            if len(table.rows) > 0:
                # Mostrar primera fila
                first_row_text = " | ".join([cell.text[:20] for cell in table.rows[0].cells])
                print(f"  First row: {first_row_text}")
                
                # Verificar si es la tabla de dispositivos
                if "CANT" in first_row_text or "EQUIPO" in first_row_text:
                    print(f"  [OK] This looks like the devices table!")
                    print(f"  Rows: {len(table.rows)}")
                    for j, row in enumerate(table.rows[:5]):  # Mostrar primeras 5 filas
                        row_text = " | ".join([cell.text[:15] for cell in row.cells])
                        print(f"    Row {j}: {row_text}")
    else:
        print("\n[ERROR] No tables found in document!")
        print("The {{TABLA}} placeholder was not replaced with a table")
    
    # Buscar si queda el placeholder sin reemplazar
    for para in doc.paragraphs:
        if "{{TABLA}}" in para.text or "{{ TABLA }}" in para.text:
            print(f"\n[ERROR] Found unreplaced placeholder: {para.text}")
    
    print("\n" + "="*80)
    print("TEST COMPLETED")
    print("="*80)
    
finally:
    db.close()
