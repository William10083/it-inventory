"""
Test actualizado para verificar lógica de negocio de tabla
"""
import sys
import os
sys.path.append(r'C:\Users\wvilca\Downloads\it_inventory 2\it_inventory\backend')

from database import SessionLocal
from models import DocumentTemplate
from pdf_generator import generate_batch_acta
from docx import Document

db = SessionLocal()

try:
    # Obtener template ID 8
    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == 8).first()
    
    print("="*80)
    print("TESTING BUSINESS LOGIC TABLE")
    print("="*80)
    
    # Datos de prueba complejos
    devices_info = [
        {
            'type': 'KEYBOARD_MOUSE_KIT',
            'brand': 'HP',
            'model': 'HSA-A005K / HSA-A011M',
            'serial': '7CH43452CM',
            'hostname': '-'
        },
        {
            'type': 'BACKPACK',
            'brand': 'HP',
            'model': 'TRAVEL 18L',
            'serial': 'BACKPACK-36',
            'hostname': '-'
        },
        {
            'type': 'MONITOR',
            'brand': 'LENOVO',
            'model': 'THINKVISION E24-30',
            'serial': 'VNABFZ0H',
            'hostname': 'TAM-MN001'
        },
        {
            'type': 'HEADPHONES',
            'brand': 'HP',
            'model': 'G2',
            'serial': '7CH4290H62',
            'hostname': '-'
        },
        {
            'type': 'LAPTOP',
            'brand': 'HP',
            'model': 'PROBOOK 440 G11',
            'serial': '5CD5401M1L',
            'hostname': 'TAM-LP001'
        },
         {
            'type': 'CHARGER',
            'brand': 'HP',
            'model': 'TPN-DA15',
            'serial': 'CHARGER-HP-225',
            'hostname': '-'
        }
    ]
    
    test_acta_path = generate_batch_acta(
        assignment_id=9999,
        employee_name="TEST BUSINESS LOGIC",
        devices_info=devices_info,
        employee_dni="12345678",
        employee_company="TRANSTOTAL",
        template_path=template.file_path,
        template=template
    )
    
    print(f"\n[OK] Test acta generated at: {test_acta_path}")
    
    # Verificar el contenido
    doc = Document(test_acta_path)
    
    print("\n" + "="*80)
    print("VERIFYING TABLE ROWS")
    print("="*80)
    
    # Buscar la tabla de 7 columnas
    target_table = None
    for table in doc.tables:
        # Buscar recursivamente
        def find_7col_table(t):
            if len(t.columns) == 7:
                return t
            for row in t.rows:
                for cell in row.cells:
                    for nested in cell.tables:
                        res = find_7col_table(nested)
                        if res: return res
            return None
        
        target_table = find_7col_table(table)
        if target_table: break
        
        # Check top level
        if len(table.columns) == 7:
            target_table = table
            break
            
    if target_table:
        print(f"Found table with {len(target_table.rows)} rows\n")
        headers = [c.text for c in target_table.rows[0].cells]
        print(f"Headers: {headers}")
        print("-" * 100)
        
        for i, row in enumerate(target_table.rows[1:], 1):
            cells = [c.text.strip() for c in row.cells]
            print(f"Row {i}: {cells}")
    else:
        print("[ERROR] Table with 7 columns not found!")

finally:
    db.close()
