"""
Test actualizado para verificar formato final y lógica de estado/mochila
"""
import sys
import os
sys.path.append(r'C:\Users\wvilca\Downloads\it_inventory 2\it_inventory\backend')

from database import SessionLocal
from models import DocumentTemplate
from pdf_generator import generate_batch_acta
from docx import Document

db = SessionLocal()

try:
    # Obtener template ID 8
    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == 8).first()
    
    print("="*80)
    print("TESTING FINAL FORMAT LOGIC")
    print("="*80)
    
    # Datos de prueba con status explícito (simulando lo que envía assignments.py)
    devices_info = [
        {
            'type': 'MONITOR',
            'brand': 'LENOVO',
            'model': 'THINKVISION',
            'serial': 'MN123',
            'hostname': 'TAM-MN001',
            'status': 'USADO'  # Monitor usado -> Cables deben ser usados
        },
        {
            'type': 'BACKPACK',
            'brand': 'HP',
            'model': 'TRAVEL',
            'serial': 'SHOULD-BE-REMOVED', # Esto debería reemplazarse por '-'
            'hostname': '-',
            'status': 'NUEVO'
        }
    ]
    
    test_acta_path = generate_batch_acta(
        assignment_id=9999,
        employee_name="TEST FINAL",
        devices_info=devices_info,
        employee_dni="12345678",
        employee_company="TRANSTOTAL",
        template_path=template.file_path,
        template=template
    )
    
    print(f"\n[OK] Test acta generated at: {test_acta_path}")
    
    # Verificar el contenido
    doc = Document(test_acta_path)
    
    # Buscar tabla recursivamente
    target_table = None
    
    def find_7col_table(tables):
        for table in tables:
            if len(table.columns) == 7:
                return table
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        found = find_7col_table(cell.tables)
                        if found: return found
        return None
    
    target_table = find_7col_table(doc.tables)
            
    if target_table:
        print(f"\nVerifying Table Content ({len(target_table.rows)} rows):")
        for i, row in enumerate(target_table.rows[1:], 1):
            cells = [c.text.strip() for c in row.cells]
            print(f"Row {i}: {cells}")
            
            # Verificar Mochila
            if "MOCHILA" in cells[1]:
                if cells[5] == '-':
                    print("  [PASS] Backpack serial is '-'")
                else:
                    print(f"  [FAIL] Backpack serial is '{cells[5]}'")
            
            # Verificar Cables
            if "CABLE" in cells[1]:
                if cells[2] == 'USADO': # Heredado del monitor
                     print("  [PASS] Cable status inherited 'USADO'")
                else:
                     print(f"  [FAIL] Cable status is '{cells[2]}', expected 'USADO'")
        
        # Verificar tamaño de fuente (muestreo)
        print("\nVerifying Font Size (Sampling):")
        try:
            # Check first data cell
            first_cell = target_table.rows[1].cells[0]
            run = first_cell.paragraphs[0].runs[0]
            if run.font.size and run.font.size.pt == 8.0:
                 print(f"  [PASS] Font size is {run.font.size.pt}")
            else:
                 print(f"  [WARN] Font size is {run.font.size.pt if run.font.size else 'None'} (Might be implicit or default)")
        except Exception as e:
            print(f"  [WARN] Could not verify font size directly: {e}")

    else:
        print("[ERROR] Table not found!")

finally:
    db.close()
