import sys
import os
import shutil

# Add backend to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from pdf_generator import generate_batch_acta, replace_placeholder_with_image
from docx import Document

def create_dummy_image(filename="dummy_test_image.png"):
    from PIL import Image, ImageDraw
    
    img = Image.new('RGB', (200, 200), color = 'red')
    d = ImageDraw.Draw(img)
    d.text((10,10), "TEST IMAGE", fill=(255,255,0))
    img.save(filename)
    return filename

def test_image_insertion():
    print("--- Probando Inserción de Imágenes en PDF ---")
    
    # 1. Crear imagen de prueba
    try:
        from PIL import Image
    except ImportError:
        print("[SKIP] PIL (Pillow) no instalado, no se puede generar imagen de prueba.")
        # Intentar usar una imagen existente si hay
        return

    image_path = create_dummy_image()
    print(f"[OK] Imagen de prueba creada: {image_path}")
    
    # 2. Crear template temporal con placeholders
    template_path = "temp_image_template.docx"
    doc = Document()
    doc.add_paragraph("Acta de Prueba de Imágenes")
    doc.add_paragraph("Foto del Dispositivo:")
    doc.add_paragraph("{{DEVICE_IMAGE_PATH}}")
    doc.add_paragraph("Foto de Serie:")
    doc.add_paragraph("{{SERIAL_IMAGE_PATH}}")
    # Agregar tabla dummy para evitar IndexError en logica legacy del generador
    doc.add_table(rows=1, cols=1)
    doc.save(template_path)
    
    # 3. Datos de prueba
    employee_data = {
        "name": "TEST USER",
        "dni": "12345678",
        "observations": "Test de imágenes",
        "decommission_data": {
            "device_image_path": os.path.abspath(image_path),
            "serial_image_path": os.path.abspath(image_path)
        }
    }
    
    # 4. Generar PDF (DOCX)
    try:
        # Nota: pasamos un objeto template simulado o None para que use el replace manual
        # En generate_batch_acta, si no hay template object, usa legacy placeholders
        # Pero nuestra logica de imagenes funciona con el diccionario 'placeholders' que se construye
        # Si pasamos template=None, usa el legacy dict
        # Debemos asegurarnos que 'DEVICE_IMAGE_PATH' sea agregado al legacy dict o procesado aparte
        # En la implementación actual, 'DEVICE_IMAGE_PATH' se agrega SI está en employee_data['decommission_data']
        # PERO... generate_batch_acta construye 'system_values' solo si hay template object?
        # NO, revisando el codigo:
        # Si no hay template object, builds legacy dict
        # Y luego busca 'DEVICE_IMAGE_PATH' en ese dict.
        # PERO legacy dict NO incluye 'DEVICE_IMAGE_PATH' por defecto!
        # Necesitamos simular un objeto template para que entre en build_placeholders_from_template
        
        class MockTemplate:
            variables = [
                {"name": "DEVICE_IMAGE_PATH", "map_to": "DEVICE_IMAGE_PATH"},
                {"name": "SERIAL_IMAGE_PATH", "map_to": "SERIAL_IMAGE_PATH"}
            ]
            template_type = "ACTA_BAJA"
            
        generated_path = generate_batch_acta(
            assignment_id=9999,
            employee_name="TEST USER",
            devices_info=[{"type": "LAPTOP", "brand": "HP", "model": "TEST", "inventory_code": "INV-001"}],
            template_path=template_path,
            template=MockTemplate(),
            decommission_data=employee_data['decommission_data']
        )
        
        print(f"[OK] Documento generado en: {generated_path}")
        
        # 5. Verificación (Manual o inspección de estructura)
        # Es difícil verificar imagen incrustada programáticamente fácil sin unzip
        print(f"[INFO] Por favor abrir {generated_path} y verificar que aparezcan cuadrados rojos.")
        
    except Exception as e:
        print(f"[X] Error generando PDF: {e}")
        import traceback
        traceback.print_exc()
        
    # Cleanup
    if os.path.exists(image_path):
        os.remove(image_path)
    if os.path.exists(template_path):
        os.remove(template_path)

if __name__ == "__main__":
    test_image_insertion()
