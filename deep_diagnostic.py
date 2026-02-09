import sys
import os
import json

# Ajustar path
sys.path.append(os.path.join(os.getcwd(), 'backend'))

from database import SessionLocal
import models
from docx import Document

def check_everything():
    db = SessionLocal()
    try:
        # 1. Check ALL templates of type ACTA_BAJA
        templates = db.query(models.DocumentTemplate).filter(models.DocumentTemplate.template_type == "ACTA_BAJA").all()
        print(f"Templates of type ACTA_BAJA in DB: {len(templates)}")
        for t in templates:
            print(f"ID: {t.id}, Name: {t.name}, Is Active: {t.is_active}, Path: {t.file_path}")
            abs_path = os.path.abspath(os.path.join(os.getcwd(), 'backend', t.file_path))
            if os.path.exists(abs_path):
                doc = Document(abs_path)
                first_text = ""
                if doc.paragraphs:
                    first_text = doc.paragraphs[0].text
                print(f"  First paragraph in disk: {first_text[:100]}")
            else:
                print(f"  FILE MISSING ON DISK: {abs_path}")
        
        # 2. Check the specific decommission record ID 5
        dec = db.query(models.Decommission).filter(models.Decommission.id == 5).first()
        if dec:
            print(f"\nDecommission ID 5:")
            print(f"  Acta Path: {dec.acta_path}")
            if dec.acta_path and os.path.exists(dec.acta_path):
                doc = Document(dec.acta_path)
                print(f"  Generated file first paragraph: {doc.paragraphs[0].text if doc.paragraphs else 'EMPTY'}")
                # Check for table headers in generated file
                for i, table in enumerate(doc.tables):
                    row_text = " ".join([c.text for c in table.rows[0].cells]).upper()
                    print(f"  Table {i} header: {row_text[:100]}")
            else:
                print(f"  GENERATED FILE MISSING: {dec.acta_path}")
                
    finally:
        db.close()

if __name__ == "__main__":
    check_everything()
