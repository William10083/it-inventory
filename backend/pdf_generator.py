# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import io
from PIL import Image, ExifTags, ImageFilter, ImageOps
from copy import deepcopy


def replace_text_in_paragraph(paragraph, placeholders):
    """
    Smart replacement for text in a paragraph.
    1. Tries to replace inside individual runs first (preserves formatting).
    2. Fallback to aggressive replacement if placeholders remain (split across runs).
    """
    
    # 1. Try to replace in runs (Preserves formatting)
    # This handles cases where the placeholder is in a single run (e.g. bolded variable)
    if paragraph.runs:
        for run in paragraph.runs:
            # Check fast against run text
            if '{' in run.text:
                for key, value in placeholders.items():
                    if key in run.text:
                        # Replace in run, keeping its style
                        run.text = run.text.replace(key, str(value))
                        # print(f"DEBUG: Replaced '{key}' in run (conserving format)")

    # 2. Check if we missed any (Split runs case)
    # If text is still containing placeholders (e.g. "{{VAR" in run1 and "}}" in run2)
    # we proceed to the aggressive replacement which rewrites the paragraph.
    full_text = paragraph.text
    needs_replacement = False
    for key in placeholders:
        if key in full_text:
            needs_replacement = True
            break
    
    if not needs_replacement:
        return
    
    # 3. Fallback: Aggressive replacement (original logic)
    # Replace all placeholders in the full text
    new_text = full_text
    replaced_count = 0
    for key, value in placeholders.items():
        if key in new_text:
            new_text = new_text.replace(key, str(value))
            replaced_count += 1
            print(f"DEBUG: Replaced '{key}' with '{value}' in paragraph (aggressive)")
    
    # If text changed, clear and recreate
    if new_text != full_text:
        # Save formatting from first run if exists
        bold = False
        italic = False
        font_name = None
        font_size = None
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            bold = first_run.bold
            italic = first_run.italic
            if first_run.font:
                font_name = first_run.font.name
                font_size = first_run.font.size
        
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        
        # Add new run with replaced text
        new_run = paragraph.add_run(new_text)
        if bold is not None:
            new_run.bold = bold
        if italic is not None:
            new_run.italic = italic
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size


def replace_text_in_table(table, placeholders):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, placeholders)
            # Recursively handle nested tables
            for nested_table in cell.tables:
                replace_text_in_table(nested_table, placeholders)

def process_document_placeholders(doc, placeholders):
    """
    Applies placeholder replacement to all parts of a Document:
    Main body, Headers, Footers, and Tables.
    Also handles Text Boxes and other complex elements by scanning the XML.
    """
    # 1. Standard body replacement (maintains formatting where possible)
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, placeholders)
    for t in doc.tables:
        replace_text_in_table(t, placeholders)
    
    # 2. Headers and Footers
    for section in doc.sections:
        for p in section.header.paragraphs: replace_text_in_paragraph(p, placeholders)
        for t in section.header.tables: replace_text_in_table(t, placeholders)
        for p in section.footer.paragraphs: replace_text_in_paragraph(p, placeholders)
        for t in section.footer.tables: replace_text_in_table(t, placeholders)
    
    # 3. GLOBAL XML REPLACEMENT (Deep Scan for Text Boxes, Shapes, etc.)
    # This is a fallback but handles things like Text Box content
    from docx.oxml.ns import qn
    
    # We'll look at all <w:t> elements in all parts of the document
    def scan_xml_elements(parent_element):
        for t in parent_element.iter(qn('w:t')):
            if not t.text: continue
            new_text = t.text
            is_replaced = False
            for key, value in placeholders.items():
                if key in new_text:
                    new_text = new_text.replace(key, str(value))
                    is_replaced = True
            
            if is_replaced:
                t.text = new_text
                # print(f"DEBUG (XML): Replaced in element text: {new_text[:50]}...")
    
    # Scan main body XML
    scan_xml_elements(doc._element)
    
    # Scan headers and footers XML
    for section in doc.sections:
        scan_xml_elements(section.header._element)
        scan_xml_elements(section.footer._element)


def copy_cell_format(source_cell, target_cell):
    """Copy formatting from source cell to target cell"""
    # Copy cell properties (borders, shading, etc.) by manipulating XML directly
    if source_cell._element.tcPr is not None:
        # Remove existing tcPr if present
        if target_cell._element.tcPr is not None:
            target_cell._element.remove(target_cell._element.tcPr)
        # Add a deep copy of source tcPr
        target_cell._element.insert(0, deepcopy(source_cell._element.tcPr))
    
    # Copy paragraph formatting from first paragraph
    if source_cell.paragraphs and target_cell.paragraphs:
        source_para = source_cell.paragraphs[0]
        target_para = target_cell.paragraphs[0]
        
        # Copy paragraph properties
        if source_para._element.pPr is not None:
            if target_para._element.pPr is not None:
                target_para._element.remove(target_para._element.pPr)
            target_para._element.insert(0, deepcopy(source_para._element.pPr))


def get_optimized_image(image_path, is_serial=False):
    """
    Carga la imagen optimizada para el documento.
    Asume que la imagen ya fue procesada en el upload, pero asegura formato compatible.
    """
    try:
        img = Image.open(image_path)
        
        # Asegurar formato RGB (evitar problemas con CMYK/Alpha en Docx)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
            
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='JPEG', quality=95) # High quality for final doc
        img_byte_arr.seek(0)
        return img_byte_arr
    except Exception as e:
        print(f"ERROR: Could not load image {image_path}: {e}")
        return image_path

def get_spanish_date():
    """Generate Spanish date string"""
    months_es = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
        7: "julio", 8: "agosto", 9: "setiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    now = datetime.now()
    return f"{now.day} de {months_es[now.month]} de {now.year}"


def replace_placeholder_with_image(doc, placeholder, image_path, width_inches=2.5):
    """
    Replaces a text placeholder with an image in the document.
    Scans paragraphs and tables.
    """
    if not image_path or not os.path.exists(image_path):
        print(f"DEBUG: Image path not found or empty: {image_path}")
        return False
        
    placeholder_variants = [
        f"{{{{{placeholder}}}}}",
        f"{{{{ {placeholder} }}}}",
        f"{{{{{placeholder.lower()}}}}}",
        f"{{{{ {placeholder.lower()} }}}}",
    ]
    
    print(f"DEBUG: Attempting to replace image placeholder '{placeholder}' with image: {image_path}")
    
    replaced = False

    def process_paragraph(paragraph):
        nonlocal replaced
        for variant in placeholder_variants:
            if variant in paragraph.text:
                # Found it! Clear text and add image
                if paragraph.text.strip() == variant or True: # Aggressive replacement
                     # Clear paragraph content
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # Centrar el párrafo y asegurar que el espaciado no corte la imagen
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing = 1.0 # Espaciado sencillo (evita cortes si era fijo)
                    paragraph.paragraph_format.space_after = Pt(2) # Pequeño margen
                    paragraph.paragraph_format.space_before = Pt(2)
                    
                    # Add image run
                    run = paragraph.add_run()
                    try:
                        # Usar imagen optimizada (corrección de rotación)
                        is_serial_placeholder = "SERIAL" in placeholder.upper()
                        optimized_img = get_optimized_image(image_path, is_serial=is_serial_placeholder)
                        
                        run.add_picture(optimized_img, width=Inches(width_inches))
                        print(f"DEBUG: Inserted optimized image for '{variant}'")
                        replaced = True
                        return True
                    except Exception as e:
                        print(f"ERROR: Failed to insert image: {e}")
                        return False
        return False

    # 1. Scan Paragraphs
    for p in doc.paragraphs:
        if process_paragraph(p):
            return True
            
    # 2. Scan Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if process_paragraph(p):
                        return True
                        
    return replaced


def insert_devices_table_at_placeholder(doc, devices_info, template=None, placeholder_text='{{TABLA}}'):
    """
    Busca el placeholder de tabla e inserta una tabla de dispositivos en esa posición.
    
    Args:
        doc: Documento Document de python-docx
        devices_info: Lista de dispositivos a insertar en la tabla
        template: Objeto DocumentTemplate con mapeo de variables (opcional)
        placeholder_text: Texto del placeholder a buscar como fallback (default: '{{TABLA}}')
    
    Returns:
        True si se insertó la tabla, False si no se encontró el placeholder
    """
    import json
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 1. Identificar el placeholder correcto desde el mapeo del template
    actual_placeholder = None
    
    if template and template.variables:
        try:
            # Parsear variables JSON
            if isinstance(template.variables, str):
                variables = json.loads(template.variables)
            else:
                variables = template.variables
            
            # Variables de tabla que buscamos
            table_variables = [
                'DEVICE_TABLE',
                'ASSIGNMENT_COMPUTER_TABLE',
                'MOBILE_DEVICES_TABLE',
                'ASSIGNMENT_MOBILE_TABLE',
                'RETURNED_DEVICES_TABLE',
                'RETURN_COMPUTER_TABLE',
                'RETURN_MOBILE_TABLE',
                'DECOMMISSION_TABLE',
                'SALE_TABLE'
            ]
            
            # Buscar qué placeholder está mapeado a una variable de tabla
            for var in variables:
                placeholder_name = var.get('name', '')
                map_to = var.get('map_to', '')
                
                if map_to in table_variables:
                    actual_placeholder = placeholder_name
                    print(f"DEBUG: Found table placeholder mapping: {placeholder_name} -> {map_to}")
                    break
        
        except Exception as e:
            print(f"DEBUG: Error parsing template variables: {e}")
    
    # 2. Si no se encontró en el mapeo, usar el fallback
    if not actual_placeholder:
        actual_placeholder = placeholder_text.replace('{{', '').replace('}}', '').strip()
        print(f"DEBUG: Using fallback placeholder: {actual_placeholder}")
    
    # 3. Buscar el párrafo con el placeholder (con variantes de espacios)
    tabla_paragraph = None
    tabla_para_element = None
    tabla_cell = None  # Para saber si está en una celda
    
    # Generar variantes del placeholder
    placeholder_variants = [
        f"{{{{{actual_placeholder}}}}}",
        f"{{{{ {actual_placeholder} }}}}",
        f"{{{{{actual_placeholder.lower()}}}}}",
        f"{{{{ {actual_placeholder.lower()} }}}}",
    ]
    
    # Función recursiva para buscar en tablas anidadas
    def search_in_tables(tables, level=0):
        nonlocal tabla_paragraph, tabla_para_element, tabla_cell
        
        print(f"DEBUG: Searching in {len(tables)} tables at level {level}")
        
        for t_idx, table in enumerate(tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    # Buscar en párrafos de la celda
                    for para in cell.paragraphs:
                        # Normalizar texto para búsqueda (eliminar espacios extraños)
                        text = para.text.strip()
                        for variant in placeholder_variants:
                            # Búsqueda más flexible
                            if variant in text or actual_placeholder in text: # Fallback a buscar solo el nombre
                                if actual_placeholder in text and not variant in text:
                                     print(f"DEBUG: Found partial match '{actual_placeholder}' in '{text}'")
                                
                                # Verificar match exacto o muy cercano
                                if variant in text:
                                    tabla_paragraph = para
                                    tabla_para_element = para._element
                                    tabla_cell = cell
                                    print(f"DEBUG: Found placeholder '{variant}' in table cell (level {level}): {para.text[:50]}")
                                    return True
                    
                    # Buscar recursivamente en tablas anidadas
                    if cell.tables:
                        if search_in_tables(cell.tables, level + 1):
                            return True
        return False
    
    # Buscar en párrafos principales
    for para in doc.paragraphs:
        for variant in placeholder_variants:
            if variant in para.text:
                tabla_paragraph = para
                tabla_para_element = para._element
                print(f"DEBUG: Found placeholder '{variant}' in main paragraph: {para.text[:50]}")
                break
        if tabla_paragraph:
            break
    
    # Si no se encontró en párrafos principales, buscar en tablas (incluyendo anidadas)
    if not tabla_paragraph:
        print(f"DEBUG: Placeholder not found in main paragraphs, searching in tables (including nested)...")
        search_in_tables(doc.tables)
    
    if not tabla_paragraph:
        print(f"DEBUG: Placeholder not found anywhere. Tried variants: {placeholder_variants}")
        return False
    
    # Identificar si es una tabla de venta basado en el placeholder o variable mapeada
    is_sales_table = 'SALE_TABLE' in (actual_placeholder or '') or 'TABLA_VENTA' in (actual_placeholder or '')
    
    # Crear tabla de dispositivos
    # Tabla Venta: 7 columnas (CANT., EQUIPO, ESTADO, MARCA, MODELO, NUMERO SERIE, CÓDIGO DE INVENTARIO)
    # Tabla Standard: 8 columnas (incluye HOSTNAME)
    
    num_cols = 7 if is_sales_table else 8
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Configurar encabezados
    hdr_cells = table.rows[0].cells
    
    if is_sales_table:
        headers = [
            'CANT.', 'EQUIPO', 'ESTADO', 'MARCA', 'MODELO', 'NUMERO\nSERIE', 'CÓDIGO DE\nINVENTARIO'
        ]
        header_bg_color = "FF0000" # Rojo
    else:
        headers = [
            'CANT.', 'EQUIPO', 'ESTADO', 'MARCA', 'MODELO', 'SERIE', 'HOSTNAME', 'CÓDIGO DE\nINVENTARIO'
        ]
        header_bg_color = "FF0000" # Usar rojo también por consistencia o mantener default? El usuario pidió rojo para venta.
        # Mantener rojo ya que parece ser el estilo deseado
    
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
    
    # Función helper para color de fondo
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    def set_cell_background(cell, color_hex):
        """Set cell background color."""
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Aplicar formato a encabezados
    try:
        from docx.shared import RGBColor
    except ImportError:
        RGBColor = None

    for cell in hdr_cells:
        # Fondo header
        set_cell_background(cell, header_bg_color)
        
        # Texto blanco, bold y centrado
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                if RGBColor:
                    run.font.color.rgb = RGBColor(255, 255, 255)

    # --- PROCESAMIENTO DE DISPOSITIVOS ---
    processed_devices = []
    
    # 1. Transformación y Desglose
    for dev in devices_info:
        dtype = (dev.get('type', '') or '').upper()
        brand = (dev.get('brand', '') or '').upper()
        model = (dev.get('model', '') or '').upper()
        serial = (dev.get('serial', '') or dev.get('imei', '') or '-').upper()
        hostname = (dev.get('hostname', '') or '').upper()
        inventory_code = (dev.get('inventory_code', '') or '').upper()
        status = (dev.get('status', 'USADO') or 'USADO').upper()
        
        # KEYBOARD_MOUSE_KIT -> Separar en TECLADO y MOUSE
        if dtype == 'KEYBOARD_MOUSE_KIT':
            # Intentar separar modelos "HSA-A005K / HSA-A011M"
            kb_model = model
            ms_model = model
            if '/' in model:
                parts = model.split('/')
                kb_model = parts[0].strip()
                ms_model = parts[1].strip() if len(parts) > 1 else kb_model
            
            # TECLADO
            processed_devices.append({
                'type': 'TECLADO',
                'brand': brand,
                'model': kb_model,
                'serial': serial, # Misma serie para ambos? Usualmente sí en kit, o serie del kit.
                'hostname': '-',
                'inventory_code': '-',
                'status': status
            })
            # MOUSE
            processed_devices.append({
                'type': 'MOUSE',
                'brand': brand,
                'model': ms_model,
                'serial': serial,
                'hostname': '-',
                'inventory_code': '-',
                'status': status
            })
            continue

        # Traducciones y Renombres
        final_type = dtype
        if dtype == 'HEADPHONES':
            final_type = 'AURICULARES'
        elif dtype == 'BACKPACK':
            final_type = 'MOCHILA'
            serial = '-'  # Mochila siempre serie '-'
        elif dtype == 'CHARGER':
            final_type = 'CARGADOR DE LAPTOP' # Default name for charger
            # Regla: Ocultar series genéricas que empiezan con CHARGER-
            if serial.upper().startswith('CHARGER-'):
                serial = '-'
        elif dtype == 'LAPTOP':
             # Sales logic: "Laptop" is fine, but check current logic
             pass
        
        # Agregar dispositivo procesado
        processed_devices.append({
            'type': final_type,
            'brand': brand,
            'model': model,
            'serial': serial,
            'hostname': hostname,
            'inventory_code': inventory_code,
            'status': status
        })


        # Agregar accesorios implícitos
        if final_type == 'MONITOR':
            # Heredar estado del monitor
            cable_status = status 
            
            # Cable HDMI
            processed_devices.append({
                'type': 'CABLE HDMI',
                'brand': '-',
                'model': '-',
                'serial': '-',
                'hostname': '-',
                'inventory_code': '-',
                'status': cable_status
            })
            # Cable de Poder
            processed_devices.append({
                'type': 'CABLE DE PODER',
                'brand': '-',
                'model': '-',
                'serial': '-',
                'hostname': '-',
                'inventory_code': '-',
                'status': cable_status
            })

    
    # 2. Ordenamiento
    order_map = {
        'LAPTOP': 1,
        'CARGADOR DE LAPTOP': 2,
        'MONITOR': 3,
        'CABLE HDMI': 4,
        'CABLE DE PODER': 5,
        'TECLADO': 6,
        'MOUSE': 7,
        'AURICULARES': 8,
        'MOCHILA': 9
    }
    
    def get_order(d):
        return order_map.get(d['type'], 99)
    
    processed_devices.sort(key=get_order)

    # Helper para establecer texto con formato
    def set_cell_text(cell, text, align=None, is_header=False):
        cell.text = text
        for paragraph in cell.paragraphs:
            if align is not None:
                paragraph.alignment = align
            for run in paragraph.runs:
                if is_header:
                    run.font.bold = True
                    if RGBColor:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    run.font.size = Pt(8)

    # Agregar filas de dispositivos processed
    for dev in processed_devices:
        row_cells = table.add_row().cells
        
        # Centrar contenido verticalmente
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        set_cell_text(row_cells[0], '1', WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row_cells[1], dev['type'])
        set_cell_text(row_cells[2], dev['status'], WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row_cells[3], dev['brand'])
        set_cell_text(row_cells[4], dev['model'])
        set_cell_text(row_cells[5], dev['serial'])
        
        if is_sales_table:
            # Columna 6: Código Inventario (en index 6)
            inventory_value = 'N/A'
            if dev.get('inventory_code') and dev['inventory_code'] != '-' and dev['inventory_code'] != '':
                inventory_value = dev['inventory_code']
            elif dev['type'] == 'LAPTOP' and dev.get('hostname'): # A veces el hostname es el codigo?
                 pass # User asked for Inventory Code explicitly
            
            # Si no hay codigo, intentamos poner '-' o lo que corresponda
            if inventory_value == 'N/A': inventory_value = '-'

            set_cell_text(row_cells[6], inventory_value)
            
        else:
            # Columna 6: Hostname (mostrar N/A si no tiene)
            hostname_value = 'N/A'
            if dev.get('hostname') and dev['hostname'] != '-' and dev['hostname'] != '':
                hostname_value = dev['hostname']
            set_cell_text(row_cells[6], hostname_value)
            
            # Columna 7: Código de Inventario (mostrar N/A si no tiene)
            inventory_value = 'N/A'
            if dev.get('inventory_code') and dev['inventory_code'] != '-' and dev['inventory_code'] != '':
                inventory_value = dev['inventory_code']
            set_cell_text(row_cells[7], inventory_value)

    
    # Insertar la tabla
    table_element = table._element
    
    if tabla_para_element is not None:
        print(f"DEBUG: Inserting table before placeholder paragraph")
        # Insertar tabla inmediatamente antes del párrafo que contiene el placeholder
        tabla_para_element.addprevious(table_element)
        
        # AGREGAR UN PÁRRAFO VACÍO PARA SEPARACIÓN
        # Esto asegura que haya un espacio entre la tabla y el texto siguiente
        new_p = OxmlElement('w:p')
        tabla_para_element.addprevious(new_p)
        
        # Eliminar el párrafo que contenía el placeholder
        parent = tabla_para_element.getparent()
        if parent is not None:
            parent.remove(tabla_para_element)
    else:
        print("ERROR: Could not find anchor paragraph element to insert table.")
    
    print(f"DEBUG: Inserted devices table with {len(devices_info)} devices")
    return True


def categorize_devices(devices_info):
    """
    Separate devices into computer equipment and mobile equipment.
    Returns: (computer_devices, mobile_devices)
    """
    computer_devices = []
    mobile_devices = []
    
    for device in devices_info:
        device_type = device.get('type', '').lower()
        if device_type in ['mobile', 'chip', 'celular', 'charger', 'cargador']:
            mobile_devices.append(device)
        else:
            computer_devices.append(device)
    
    return computer_devices, mobile_devices


def build_placeholders_from_template(template, employee_data, devices_info=None):
    """
    Construye placeholders dinámicamente desde el mapeo del template.
    
    Args:
        template: Objeto DocumentTemplate con el mapeo de variables (puede ser None)
        employee_data: Dict con datos del empleado {name, dni, company, location, observations, template_type}
        devices_info: Lista de dispositivos (opcional, para futuro uso)
    
    Returns:
        Dict de placeholders {placeholder: valor}
    """
    import json
    
    placeholders = {}
    
    # Si no hay template o no tiene variables, retornar vacío
    if not template or not template.variables:
        return placeholders
    
    # Parse variables JSON
    try:
        if isinstance(template.variables, str):
            variables = json.loads(template.variables)
        else:
            variables = template.variables
    except Exception as e:
        print(f"ERROR: Could not parse template variables: {e}")
        return placeholders
    
    # Determinar tipo de template para observaciones por defecto
    template_type = template.template_type if template else employee_data.get('template_type', '')
    
    # Observaciones por defecto según tipo de template
    if 'ASSIGNMENT' in template_type:
        default_observations = "Se hace entrega de los equipos asignados en óptimas condiciones para el desempeño de sus funciones."
    elif 'RETURN' in template_type:
        default_observations = "Se recibe equipos completos en buen estado de conservación y funcionamiento."
    elif 'ACTA_BAJA' in template_type:
        default_observations = "Se procede a la baja del activo del inventario por motivo justificado."
    else:
        default_observations = ""
    
    # Mapeo de variables del sistema a valores
    system_values = {
        'EMPLOYEE_NAME': employee_data.get('name', '').upper(),
        'EMPLOYEE_DNI': employee_data.get('dni', ''),
        'EMPLOYEE_COMPANY': employee_data.get('company', '').upper(),
        'EMPLOYEE_LOCATION': employee_data.get('location', 'CALLAO').upper(),
        'CURRENT_DATE_LONG': get_spanish_date(),
        'CURRENT_DATE': datetime.now().strftime('%d/%m/%Y'),
        
        # Observaciones diferenciadas
        'ACTA_OBSERVATIONS': employee_data.get('observations', default_observations),
        'ASSIGNMENT_OBSERVATIONS': "Se hace entrega de los equipos asignados en óptimas condiciones para el desempeño de sus funciones.",
        'RETURN_OBSERVATIONS': "Se recibe equipos completos en buen estado de conservación y funcionamiento.",
        
        # Variables de tabla (4 tipos)
        'ASSIGNMENT_COMPUTER_TABLE': '[TABLA_ENTREGA_COMPUTO]',  # Placeholder, se llena dinámicamente
        'ASSIGNMENT_MOBILE_TABLE': '[TABLA_ENTREGA_CELULAR]',
        'RETURN_COMPUTER_TABLE': '[TABLA_DEVOLUCION_COMPUTO]',
        'RETURN_MOBILE_TABLE': '[TABLA_DEVOLUCION_CELULAR]',
        'RETURN_MOBILE_TABLE': '[TABLA_DEVOLUCION_CELULAR]',
        'DECOMMISSION_TABLE': '[TABLA_BAJA]',
        'SALE_TABLE': '[TABLA_VENTA]',
        
        # Aliases legacy para compatibilidad
        'DEVICE_TABLE': '[TABLA_DISPOSITIVOS]',
        'MOBILE_DEVICES_TABLE': '[TABLA_CELULARES]',
        'RETURNED_DEVICES_TABLE': '[TABLA_DEVOLUCION]',

        
        # Specific aliases
        'DECOMMISSION_OBSERVATIONS': employee_data.get('observations', default_observations),
    }
    
    # Add decommission-specific variables if provided
    if employee_data.get('decommission_data'):
        decom_data = employee_data['decommission_data']
        current_year = datetime.now().year
        
        # Fabrication year
        fabrication_year = decom_data.get('fabrication_year')
        system_values['FABRICATION_YEAR'] = str(fabrication_year) if fabrication_year else 'N/A'
        
        # Usage time (calculated)
        if fabrication_year:
            usage_time = current_year - fabrication_year
            system_values['USAGE_TIME'] = f"{usage_time} años" if usage_time > 0 else "N/A"
        else:
            system_values['USAGE_TIME'] = 'N/A'
        
        # Purchase reason
        purchase_reason = decom_data.get('purchase_reason')
        system_values['PURCHASE_REASON'] = purchase_reason if purchase_reason else 'N/A'
        
        # Device and serial image paths (for future use)
        system_values['DEVICE_IMAGE_PATH'] = decom_data.get('device_image_path', '')
        system_values['SERIAL_IMAGE_PATH'] = decom_data.get('serial_image_path', '')
    
    # Current date components (always available)
    now = datetime.now()
    system_values['CURRENT_YEAR'] = str(now.year)
    
    # Spanish month names in uppercase
    months_es_upper = {
        1: "ENERO", 2: "FEBRERO", 3: "MARZO", 4: "ABRIL", 5: "MAYO", 6: "JUNIO",
        7: "JULIO", 8: "AGOSTO", 9: "SETIEMBRE", 10: "OCTUBRE", 11: "NOVIEMBRE", 12: "DICIEMBRE"
    }
    system_values['CURRENT_MONTH'] = months_es_upper[now.month]

    # Device Info (Main Device - for single-device templates like Decommission)
    if devices_info and len(devices_info) > 0:
        main_device = devices_info[0]
        inventory_code = main_device.get('inventory_code')
        system_values['DEVICE_INVENTORY_CODE'] = inventory_code if inventory_code and inventory_code != '-' else 'N/A'
        system_values['DEVICE_MODEL'] = (main_device.get('model') or '').upper()
        system_values['DEVICE_BRAND'] = (main_device.get('brand') or '').upper()
        system_values['DEVICE_SERIAL'] = (main_device.get('serial') or main_device.get('imei') or '').upper()
        system_values['DEVICE_TYPE'] = (main_device.get('type') or '').upper()
        
        # Mapping AURICULARES and MOCHILA
        if system_values['DEVICE_TYPE'] == 'HEADPHONES': system_values['DEVICE_TYPE'] = 'AURICULARES'
        if system_values['DEVICE_TYPE'] == 'BACKPACK': system_values['DEVICE_TYPE'] = 'MOCHILA'
        if system_values['DEVICE_TYPE'] == 'CHARGER': system_values['DEVICE_TYPE'] = 'CARGADOR'
    
    # Lista de variables de tabla que NO deben ser reemplazadas como texto
    # Estas serán manejadas por insert_devices_table_at_placeholder()
    table_variables = [
        'DEVICE_TABLE',
        'ASSIGNMENT_COMPUTER_TABLE',
        'MOBILE_DEVICES_TABLE',
        'ASSIGNMENT_MOBILE_TABLE',
        'RETURNED_DEVICES_TABLE',
        'RETURN_COMPUTER_TABLE',
        'RETURN_MOBILE_TABLE',
        'RETURN_MOBILE_TABLE',
        'DECOMMISSION_TABLE',
        'SALE_TABLE'
    ]
    
    # Construir placeholders desde el mapeo
    for var in variables:
        placeholder_name = var.get('name', '')
        system_var = var.get('map_to', '')
        
        if not placeholder_name or not system_var:
            continue
        
        # SKIP variables de tabla - serán manejadas por insert_devices_table_at_placeholder()
        if system_var in table_variables:
            print(f"DEBUG: Skipping table variable {placeholder_name} -> {system_var} (will be handled by table insertion)")
            continue
        
        if system_var in system_values:
            # Agregar con {{}} y variantes con espacios
            placeholders[f"{{{{{placeholder_name}}}}}"] = system_values[system_var]
            placeholders[f"{{{{ {placeholder_name} }}}}"] = system_values[system_var]
            # También agregar variantes en minúsculas y capitalize
            placeholders[f"{{{{{placeholder_name.lower()}}}}}"] = system_values[system_var]
            placeholders[f"{{{{ {placeholder_name.lower()} }}}}"] = system_values[system_var]
    
    print(f"DEBUG: Built {len(placeholders)} placeholders from template mapping (excluding table variables)")
    return placeholders


def generate_batch_acta(assignment_id: int, employee_name: str, devices_info: list, employee_dni: str = "", employee_company: str = "", template_path: str = None, template=None, acta_observations: str = None, decommission_data: dict = None):
    # Generates a PDF for a batch of devices with dynamic table rows.
    # Uses the provided template_path or defaults to acta_template.docx.
    # Args:
    #     decommission_data: Optional dict with decommission-specific data: (fabrication_year, purchase_reason, device_image_path, serial_image_path)
    # Setup Output
    filename = f"acta_{assignment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_dir = os.path.join(os.path.dirname(__file__), "generated_pdfs")
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)

    # Observaciones por defecto si están vacías
    if not acta_observations or acta_observations.strip() == "-" or acta_observations.strip() == "":
        acta_observations = "Entrega de equipos de cómputo y accesorios para el desempeño de funciones laborales."
    
    # Load Template
    is_decommission = decommission_data is not None
    
    if template and hasattr(template, 'file_path'):
        template_path = template.file_path
        print(f"DEBUG: Using template path from DB object: {template_path}")

    if template_path:
        # Resolver ruta absoluta siempre para evitar errores de contexto
        if not os.path.isabs(template_path):
            # Probar relativa al directorio backend
            abs_path = os.path.abspath(os.path.join(os.path.dirname(__file__), template_path))
            if os.path.exists(abs_path):
                template_path = abs_path
            else:
                # Probar relativa a la raíz si no se encontró
                abs_path_root = os.path.abspath(os.path.join(os.getcwd(), template_path))
                if os.path.exists(abs_path_root):
                    template_path = abs_path_root
        
        # Verificar si existe antes de proceder
        if not os.path.exists(template_path):
            print(f"DEBUG: Template NOT FOUND at: {template_path}")
            template_path = None

    # Lógica de Fallback inteligente
    if not template_path:
        if is_decommission:
            # Para bajas, no queremos usar el acta de entrega por error
            # Intentar buscar el template por defecto de baja en la ruta estándar
            default_baja = os.path.join(os.path.dirname(__file__), "templates", "INFORME_DE_BAJA.docx")
            if os.path.exists(default_baja):
                template_path = default_baja
                print(f"DEBUG: Using default baja template: {template_path}")
            else:
                # Si no hay nada, usar el fallback general pero loguear el error
                template_path = os.path.join(os.path.dirname(__file__), "templates", "acta_template.docx")
                print(f"WARNING: No decommission template found. USING FALLBACK: {template_path}")
        else:
            # Lógica normal para asignaciones
            template_path = os.path.join(os.path.dirname(__file__), "templates", "acta_template.docx")
            print(f"DEBUG: Using default assignment fallback template: {template_path}")
    else:
        print(f"DEBUG: SUCCESS - Loading template at: {template_path}")
    
    print(f"DEBUG: Final Template Path decided: {template_path}")
    doc = Document(template_path)
    date_str = get_spanish_date()
    
    print(f"DEBUG: Preparing placeholders. Template object present: {template is not None}")
    
    # Debug message injection removed by user request
    pass

    if template:
        print(f"DEBUG: Template Name: {template.name}, Type: {template.template_type}")

        # === PREPARE PLACEHOLDERS === (Resto del código...)


    # === PREPARE PLACEHOLDERS ===
    name_upper = employee_name.upper()
    dni_str = employee_dni if employee_dni else ""
    comp_upper = employee_company.upper() if employee_company else ""
    
    # Determinar tipo de template para el mapeo
    t_type = 'ASSIGNMENT_COMPUTER'
    if is_decommission: t_type = 'ACTA_BAJA'
    if template: t_type = template.template_type

    # Preparar datos del empleado para el mapeo dinámico
    employee_data = {
        'name': employee_name,
        'dni': dni_str,
        'company': employee_company,
        'location': 'CALLAO',
        'observations': acta_observations,
        'template_type': t_type,
        'decommission_data': decommission_data
    }
    
    # Construir placeholders dinámicamente desde el template
    if template:
        placeholders = build_placeholders_from_template(template, employee_data, devices_info)
        print(f"DEBUG: Using dynamic placeholders from template (count: {len(placeholders)})")
    else:
        # Fallback: usar placeholders legacy si no hay template
        print("DEBUG: No template provided, using legacy placeholders")
        placeholders = {
            "«NOMBRE»": name_upper,
            "{{NOMBRE}}": name_upper,
            "{{nombre}}": name_upper,
            "{{Nombre}}": name_upper,
            "{{ NOMBRE }}": name_upper,
            "{{ nombre }}": name_upper,
            "{{ NOMBRE_EMPLEADO}}": name_upper,
            "{{NOMBRE_EMPLEADO }}": name_upper,
            "{{ NOMBRE_EMPLEADO }}": name_upper,
            "{{NOMBRE_EMPLEADO}}": name_upper,
            "{{nombre_empleado}}": name_upper,
            "{{Nombre_Empleado}}": name_upper,
            "{{ nombre_empleado }}": name_upper,
            "{{DNI}}": dni_str,
            "{{dni}}": dni_str,
            "{{Dni}}": dni_str,
            "{{ DNI }}": dni_str,
            "{{ dni }}": dni_str,
            "{{ Dni }}": dni_str,
            "{{EMPRESA}}": comp_upper,
            "{{empresa}}": comp_upper,
            "{{Empresa}}": comp_upper,
            "{{ EMPRESA }}": comp_upper,
            "{{ empresa }}": comp_upper,
            "{{FECHA}}": date_str,
            "{{fecha}}": date_str,
            "{{Fecha}}": date_str,
            "{{ FECHA }}": date_str,
            "{{ fecha }}": date_str,
            "{{FECHA_LARGA}}": date_str,
            "{{fecha_larga}}": date_str,
            "{{Fecha_Larga}}": date_str,
            "{{ FECHA_LARGA }}": date_str,
            "{{ fecha_larga }}": date_str,
            "{{FECHALARGA}}": date_str,
            "{{ FECHALARGA }}": date_str,
            "{{SEDE}}": "CALLAO",
            "{{sede}}": "Callao",
            "{{Sede}}": "Callao",
            "{{ SEDE }}": "CALLAO",
            "{{ sede }}": "Callao",
            "{{USUARIO}}": name_upper,
            "{{ USUARIO }}": name_upper
        }
    
    # === PROCESS IMAGES ===
    # Check for image variables in template mapping
    device_image_placeholder = 'DEVICE_IMAGE_PATH' # Fallback
    serial_image_placeholder = 'SERIAL_IMAGE_PATH' # Fallback
    
    # Si tenemos un template con mapeo, buscar los nombres de los placeholders reales
    if template and template.variables:
        try:
            import json
            mapping = json.loads(template.variables) if isinstance(template.variables, str) else template.variables
            for var in mapping:
                if var.get('map_to') == 'DEVICE_IMAGE_PATH':
                    device_image_placeholder = var.get('name')
                elif var.get('map_to') == 'SERIAL_IMAGE_PATH':
                    serial_image_placeholder = var.get('name')
        except:
            pass

    # Get image paths from decommission_data
    device_img = None
    serial_img = None
    
    if decommission_data:
        device_img = decommission_data.get('device_image_path')
        serial_img = decommission_data.get('serial_image_path')
    
    def resolve_image_path(img_path):
        if not img_path: return None
        if os.path.isabs(img_path): return img_path if os.path.exists(img_path) else None
        
        possible_paths = [
            os.path.join(os.path.dirname(os.path.dirname(__file__)), img_path), # Root
            os.path.join(os.path.dirname(__file__), img_path) # Backend
        ]
        for p in possible_paths:
            if os.path.exists(p): return p
        return None

    # Replace Device Image
    final_device_img = resolve_image_path(device_img)
    if final_device_img:
        print(f"DEBUG: Replacing {device_image_placeholder} with {final_device_img}")
        replace_placeholder_with_image(doc, device_image_placeholder, final_device_img)
    
    # Replace Serial Image
    final_serial_img = resolve_image_path(serial_img)
    if final_serial_img:
        print(f"DEBUG: Replacing {serial_image_placeholder} with {final_serial_img}")
        replace_placeholder_with_image(doc, serial_image_placeholder, final_serial_img)

    
    print(f"DEBUG: Preparing placeholders for {name_upper} (DNI: {dni_str})")
    
    # Add prefixed variables for computer devices
    for dev in devices_info:
        dtype = (dev.get('type', '') or '').upper()
        # Primary variants - use 'or ""' to handle None values
        placeholders[f"{{{{SERIE_{dtype}}}}}"] = (dev.get('serial', '') or '').upper()
        placeholders[f"{{{{MARCA_{dtype}}}}}"] = (dev.get('brand', '') or '').upper()
        placeholders[f"{{{{MODELO_{dtype}}}}}"] = (dev.get('model', '') or '').upper()
        placeholders[f"{{{{HOSTNAME_{dtype}}}}}"] = (dev.get('hostname', '') or '').upper()
        # Variants with spaces
        placeholders[f"{{{{ SERIE_{dtype} }}}}"] = (dev.get('serial', '') or '').upper()
        placeholders[f"{{{{ MARCA_{dtype} }}}}"] = (dev.get('brand', '') or '').upper()
        placeholders[f"{{{{ MODELO_{dtype} }}}}"] = (dev.get('model', '') or '').upper()
        placeholders[f"{{{{ HOSTNAME_{dtype} }}}}"] = (dev.get('hostname', '') or '').upper()
    
    print(f"DEBUG: Final placeholders count: {len(placeholders)}")

    # === TRY TO INSERT TABLE AT PLACEHOLDER FIRST (BEFORE REPLACING OTHER PLACEHOLDERS) ===
    # Skip table insertion for Decommission (Baja) reports as per user request
    tabla_inserted = False
    if not is_decommission:
        tabla_inserted = insert_devices_table_at_placeholder(doc, devices_info, template=template)

    
    if tabla_inserted:
        print("DEBUG: Table inserted at {{TABLA}} placeholder")
    else:
        print("DEBUG: Table NOT inserted (placeholder not found or error)")
    
    # === REPLACE ALL OTHER PLACEHOLDERS ===
    process_document_placeholders(doc, placeholders)
    
    if tabla_inserted or is_decommission:
        # Si se insertó la tabla o es una baja (que usa placeholders individuales), guardar y retornar
        # No queremos ejecutar la lógica legada de asignaciones en una baja.
        print(f"DEBUG: Saving and returning {('DECOMMISSION' if is_decommission else 'TABLE')} document.")
        doc.save(filepath)
        return filepath

    # === LEGACY/FIXED TABLE LOGIC ===
    # Get main table for signature section
    main_table = doc.tables[0]

    # === HANDLE DATE ===
    for row in main_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "Callao," in para.text and "de" in para.text and "de 20" in para.text:
                    para.text = date_str
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # === FIND AND POPULATE DEVICES TABLE ===
    devices_table = None
    print(f"DEBUG: Searching for devices table in document with {len(doc.tables)} tables")
    
    # First, try to find in nested tables (original logic)
    for row in main_table.rows:
        for cell in row.cells:
            if cell.tables:
                print(f"DEBUG: Found {len(cell.tables)} nested table(s) in a cell")
                for nested in cell.tables:
                    if len(nested.rows) > 0:
                        first_row_text = " ".join([c.text for c in nested.rows[0].cells]).upper()
                        print(f"DEBUG: Nested table first row: {first_row_text[:100]}")
                        if "CANT" in first_row_text and "EQUIPO" in first_row_text:
                            devices_table = nested
                            print(f"DEBUG: FOUND devices table in nested! Rows: {len(nested.rows)}, Cols: {len(nested.rows[0].cells)}")
                            break
                if devices_table:
                    break
        if devices_table:
            break
    
    # If not found in nested tables, search in all document tables
    if not devices_table:
        print("DEBUG: Not found in nested tables, searching in main document tables")
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                # Check ALL rows, not just the first one (some templates have empty first rows)
                for row_idx, row in enumerate(table.rows):
                    row_text = " ".join([c.text for c in row.cells]).upper()
                    if row_idx == 0:  # Log first row for debugging
                        print(f"DEBUG: Table {table_idx} first row: {row_text[:100]}")
                    if "CANT" in row_text and "EQUIPO" in row_text:
                        devices_table = table
                        print(f"DEBUG: FOUND devices table in main tables (row {row_idx})! Rows: {len(table.rows)}, Cols: {len(table.rows[0].cells)}")
                        break
                if devices_table:
                    break
    
    if devices_table:
        print(f"DEBUG: Found devices table with {len(devices_table.rows)} rows")
        
        # Structure: Row 0 = Header, Rows 1-2 = Template data, Rows 3-4 = OBSERVACIONES
        # We need to:
        # 1. Save formatting from row 1
        # 2. Remove rows 1 and 2 (template data)
        # 3. Insert our device rows before OBSERVACIONES
        
        template_row_format = None
        if len(devices_table.rows) > 1:
            # Save the template row's element before we remove it
            template_row_format = devices_table.rows[1]
        
        # Find OBSERVACIONES row index
        obs_index = None
        for i, row in enumerate(devices_table.rows):
            if "OBSERVACIONES" in " ".join([c.text for c in row.cells]):
                obs_index = i
                break
        
        # Determine which rows to remove (template data rows, between header and observaciones)
        # In this template: row 1 (LAPTOP), row 2 (CARGADOR)
        rows_to_remove_indices = []
        for i in range(1, obs_index if obs_index else len(devices_table.rows)):
            rows_to_remove_indices.append(i)
        
        # Remove template data rows (go in reverse to maintain indices)
        for i in sorted(rows_to_remove_indices, reverse=True):
            if i < len(devices_table.rows):
                devices_table._tbl.remove(devices_table.rows[i]._tr)
        
        print(f"DEBUG: Removed {len(rows_to_remove_indices)} template rows")
        
        # Device type translations
        type_map = {
            "LAPTOP": "LAPTOP",
            "CHARGER": "CARGADOR",
            "MONITOR": "MONITOR",
            "KEYBOARD": "TECLADO",
            "MOUSE": "MOUSE",
            "KEYBOARD_MOUSE_KIT": "KIT TECLADO/MOUSE",
            "BACKPACK": "MOCHILA",
            "STAND": "SOPORTE",
            "HEADPHONES": "AURICULARES",
        }
        
        # Define device order priority
        device_order = {
            "LAPTOP": 1,
            "CHARGER": 2,
            "MONITOR": 3,
            "CABLE DE PODER": 4,  # Monitor accessory
            "CABLE HDMI": 5,      # Monitor accessory
            "KEYBOARD_MOUSE_KIT": 6,
            "KEYBOARD": 7,
            "MOUSE": 8,
            "BACKPACK": 9,
            "HEADPHONES": 10,
            "STAND": 11,
        }
        # Sort devices by priority order
        def get_order(device):
            device_type = device.get('type', '').upper()
            return device_order.get(device_type, 99)
        
        # Filter for computer equipment ONLY for this acta
        comp_devices, _ = categorize_devices(devices_info)
        sorted_devices = sorted(comp_devices, key=get_order)
        
        # Build final device list with accessories inserted after their parent devices
        final_devices = []
        for device in sorted_devices:
            final_devices.append(device)
            device_type = device.get('type', '').upper()
            
            # If this is a laptop, add charger right after (unless a charger is already in the list)
            if device_type == 'LAPTOP':
                # Check if there's already a charger in the devices list
                has_charger = any(d.get('type', '').upper() == 'CHARGER' for d in sorted_devices)
                if not has_charger:
                    final_devices.append({
                        'type': 'CARGADOR', 
                        'brand': 'HP', 
                        'model': 'TPN-DA15', 
                        'serial': '-', 
                        'hostname': '-', 
                        '_is_accessory': True
                    })
            
            # If this is a monitor, add its accessories right after
            if device_type == 'MONITOR':
                final_devices.append({'type': 'CABLE DE PODER', 'brand': '-', 'model': '-', 'serial': '-', 'hostname': '-', '_is_accessory': True})
                final_devices.append({'type': 'CABLE HDMI', 'brand': '-', 'model': '-', 'serial': '-', 'hostname': '-', '_is_accessory': True})
        
        # Find and save OBSERVACIONES row (it's the last row of the table)
        obs_text = None
        for row in devices_table.rows:
            row_text = " ".join([c.text for c in row.cells])
            if "OBSERVACIONES" in row_text:
                # Save just the first cell text (which should be "OBSERVACIONES:")
                # and the second cell which has the description
                if len(row.cells) > 0:
                    obs_text = row.cells[0].text.strip()
                    if len(row.cells) > 1 and row.cells[1].text.strip():
                        obs_description = row.cells[1].text.strip()
                    else:
                        obs_description = "- Asignación de equipo computo."
                break
        
        # Remove ALL non-header rows (template data + OBSERVACIONES)
        while len(devices_table.rows) > 1:
            devices_table._tbl.remove(devices_table.rows[1]._tr)
        
        print(f"DEBUG: Cleared all non-header rows")
        print(f"DEBUG: Will add {len(final_devices)} device rows")
        
        # Add device rows using simple add_row
        for idx, device in enumerate(final_devices):
            new_row = devices_table.add_row()
            cells = new_row.cells
            
            device_type = device.get('type', '').upper()
            device_type_display = type_map.get(device_type, device_type)
            hostname = device.get('hostname', '-')
            inventory_name = device.get('inventory_name', '-')
            
            # Populate all 8 columns: CANT, EQUIPO, ESTADO, MARCA, MODELO, SERIE, (empty), NOMBRE EQUIPO/INVENTARIO
            if len(cells) >= 1:
                cells[0].text = "1"  # CANT
            if len(cells) >= 2:
                cells[1].text = device_type_display  # EQUIPO
            if len(cells) >= 3:
                cells[2].text = "Nuevo"  # ESTADO
            if len(cells) >= 4:
                cells[3].text = (device.get('brand') or '').upper()  # MARCA
            if len(cells) >= 5:
                cells[4].text = (device.get('model') or '').upper()  # MODELO
            if len(cells) >= 6:
                cells[5].text = (device.get('serial') or '').upper()  # SERIE
            if len(cells) >= 7:
                cells[6].text = ""  # Empty column (possibly for manual notes)
            if len(cells) >= 8:
                cells[7].text = hostname if hostname else "-"  # NOMBRE EQUIPO/INVENTARIO
            
            # Apply formatting to all populated columns
            for cell_idx in range(min(8, len(cells))):
                cell = cells[cell_idx]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
            
            print(f"DEBUG: Added device row {idx}: {device_type_display} (8 columns)")
        
        # Re-add OBSERVACIONES row at the end (should span all columns)
        if obs_text:
            obs_row = devices_table.add_row()
            
            # Merge all cells in the OBSERVACIONES row to span the full width
            if len(obs_row.cells) > 1:
                # Merge from first cell to last cell
                merged_cell = obs_row.cells[0].merge(obs_row.cells[-1])
            else:
                merged_cell = obs_row.cells[0]
            
            # Set the text in the merged cell
            obs_description_text = obs_description if 'obs_description' in locals() and obs_description else "- Asignación de equipo computo."
            # obs_text already contains "OBSERVACIONES:", so just add the description
            merged_cell.text = f"OBSERVACIONES: {obs_description_text}"
            
            # Make OBSERVACIONES text bold
            for para in merged_cell.paragraphs:
                if para.runs:
                    para.runs[0].bold = True
            
            print(f"DEBUG: Re-added OBSERVACIONES row (merged across all columns)")
        
        print(f"DEBUG: Added {len(final_devices)} device rows")
        print(f"DEBUG: Table now has {len(devices_table.rows)} rows")
        for i, row in enumerate(devices_table.rows):
            row_preview = " | ".join([c.text[:15] for c in row.cells[:4]])
            print(f"  Row {i}: {row_preview}")
    else:
        print("DEBUG: Could not find devices table in template")

    doc.save(filepath)
    return filepath




def generate_mobile_acta(assignment_id: int, employee_name: str, devices_info: list, employee_dni: str = "", employee_company: str = "", template_path: str = None, template=None, acta_observations: str = None):
    # Generates acta for mobile devices using provided or default mobile template.
    # Setup Output
    filename = f"acta_celular_{assignment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_dir = os.path.join(os.path.dirname(__file__), "generated_pdfs")
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)

    # Observaciones por defecto si están vacías
    if not acta_observations or acta_observations.strip() == "-" or acta_observations.strip() == "":
        acta_observations = "Entrega de equipos de comunicación y accesorios para el desempeño de funciones laborales."

    # Load Mobile Template
    if template_path:
        # Si la ruta no es absoluta, resolverla relativa al directorio backend
        if not os.path.isabs(template_path):
            template_path = os.path.join(os.path.dirname(__file__), template_path)
            print(f"DEBUG: Resolved relative path to: {template_path}")
        
        # Verificar si existe
        if not os.path.exists(template_path):
            print(f"DEBUG: Template not found at: {template_path}")
            template_path = None
    
    # Si no hay template_path, usar fallback
    if not template_path:
        template_path = os.path.join(os.path.dirname(__file__), "templates", "acta_celular_template.docx")
        print(f"DEBUG: Using fallback mobile template: {template_path}")
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Mobile template not found at {template_path}")
    
    doc = Document(template_path)
    date_str = get_spanish_date()

    # Prepare placeholders
    name_upper = employee_name.upper()
    dni_str = employee_dni if employee_dni else ""
    comp_upper = employee_company.upper() if employee_company else "TRANSTOTAL AGENCIA MARITIMA S.A."
    
    # Preparar datos del empleado para el mapeo dinámico
    employee_data = {
        'name': employee_name,
        'dni': dni_str,
        'company': employee_company,
        'location': 'CALLAO',
        'observations': acta_observations,
        'template_type': 'ASSIGNMENT_MOBILE'
    }
    
    # Construir placeholders dinámicamente desde el template
    if template:
        placeholders = build_placeholders_from_template(template, employee_data, devices_info)
        print(f"DEBUG: Using dynamic placeholders from template (count: {len(placeholders)})")
    else:
        # Fallback: usar placeholders legacy si no hay template
        print("DEBUG: No template provided, using legacy placeholders")
        placeholders = {
            "¿NOMBRE¿": name_upper,
            "{{NOMBRE}}": name_upper,
            "{{nombre}}": name_upper,
            "{{Nombre}}": name_upper,
            "{{ NOMBRE }}": name_upper,
            "{{ nombre }}": name_upper,
            "{{NOMBRE_EMPLEADO}}": name_upper,
            "{{nombre_empleado}}": name_upper,
            "{{Nombre_Empleado}}": name_upper,
            "{{ NOMBRE_EMPLEADO }}": name_upper,
            "{{ nombre_empleado }}": name_upper,
            "{{DNI}}": dni_str,
            "{{dni}}": dni_str,
            "{{Dni}}": dni_str,
            "{{ DNI }}": dni_str,
            "{{ dni }}": dni_str,
            "{{EMPRESA}}": comp_upper,
            "{{empresa}}": comp_upper,
            "{{Empresa}}": comp_upper,
            "{{ EMPRESA }}": comp_upper,
            "{{ empresa }}": comp_upper,
            "{{FECHA}}": date_str,
            "{{fecha}}": date_str,
            "{{Fecha}}": date_str,
            "{{ FECHA }}": date_str,
            "{{ fecha }}": date_str,
            "{{FECHA_LARGA}}": date_str,
            "{{fecha_larga}}": date_str,
            "{{Fecha_Larga}}": date_str,
            "{{ FECHA_LARGA }}": date_str,
            "{{ fecha_larga }}": date_str,
            "{{SEDE}}": "CALLAO",
            "{{sede}}": "Callao",
            "{{Sede}}": "Callao",
            "{{ SEDE }}": "CALLAO",
            "{{ sede }}": "Callao",
            "{{FECHALARGA}}": date_str,
            "{{ FECHALARGA }}": date_str,
            "{{USUARIO}}": name_upper,
            "{{ USUARIO }}": name_upper
        }
    
    # Add prefixed variables for mobile devices (MOBILE, CHIP)
    for dev in devices_info:
        dtype = dev.get('type', '').upper()
        if dtype == "MOBILE": dtype = "CELULAR"
        if dtype == "CHIP": dtype = "CHIPSIM"
        
        placeholders[f"{{{{SERIE_{dtype}}}}}"] = dev.get('serial', '').strip() or dev.get('imei', '-')
        placeholders[f"{{{{MARCA_{dtype}}}}}"] = dev.get('brand', '').upper()
        placeholders[f"{{{{MODELO_{dtype}}}}}"] = dev.get('model', '').upper()
        placeholders[f"{{{{IMEI_{dtype}}}}}"] = dev.get('imei', '').upper()

    # === REPLACE IN ALL DOCUMENT PARTS ===
    process_document_placeholders(doc, placeholders)
    
    # Find devices table
    def find_devices_table(tables):
        for table in tables:
            if len(table.rows) > 0:
                first_row_text = " ".join([cell.text for cell in table.rows[0].cells]).upper()
                if "CANT" in first_row_text and "EQUIPO" in first_row_text:
                    return table
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        result = find_devices_table(cell.tables)
                        if result:
                            return result
        return None
    
    devices_table = find_devices_table(doc.tables)
    
    if devices_table:
        # Save formatting from a data row
        template_row = None
        if len(devices_table.rows) > 1:
            template_row = devices_table.rows[1]
        
        # Clear existing rows
        while len(devices_table.rows) > 1:
            devices_table._tbl.remove(devices_table.rows[-1]._tr)
        
        # Add rows for mobile devices
        for device in devices_info:
            row = devices_table.add_row()
            cells = row.cells
            
            # Copy formatting from template row if available
            if template_row:
                for i, cell in enumerate(cells):
                    if i < len(template_row.cells):
                        copy_cell_format(template_row.cells[i], cell)
            
            device_type_display = device['type'].upper()
            if device_type_display == "MOBILE":
                device_type_display = "CELULAR"
            elif device_type_display == "CHIP":
                device_type_display = "CHIP/SIM"
            
            num_cols = len(cells)
            if num_cols >= 1:
                cells[0].text = "1"  # CANT
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if num_cols >= 2:
                cells[1].text = device_type_display  # EQUIPO
            if num_cols >= 3:
                cells[2].text = "CLARO"  # OPERADOR
            if num_cols >= 4:
                cells[3].text = "Nuevo"  # ESTADO
            if num_cols >= 5:
                cells[4].text = device['brand'].upper()  # MARCA
            if num_cols >= 6:
                cells[5].text = device['model'].upper()  # MODELO
            if num_cols >= 7:
                cells[6].text = device.get('serial', '-').upper()  # SERIE
            if num_cols >= 8:
                cells[7].text = device.get('imei', '-')  # IMEI
            if num_cols >= 9:
                cells[8].text = device.get('phone_number', '-')  # TELEFONO
    
    # Handle Date
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "Callao," in p.text and "de" in p.text:
                        p.text = date_str
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(filepath)
    return filepath


from utils import doc_utils

# ... (existing imports, but remove docx if no longer needed for THIS function, though other functions use it)

# ... (keep existing helper functions not replaced by doc_utils)

def generate_return_acta_computer(termination_id, employee, assignments, observations=None, template_path=None, template_variables=None):
    # Generate return acta for computer equipment.
    # If template_path is provided, use dynamic rendering via doc_utils.
    if template_path and os.path.exists(template_path) and template_variables:
        filename = f"Acta_Recepcion_Computadora_{termination_id}_{employee.full_name.replace(' ', '_')}.docx"
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "actas")
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, filename)
        
        # Prepare Data Map
        final_vars = {}
        table_map = {}
        date_str = get_spanish_date()
        
        for placeholder, system_field in template_variables.items():
            val = ""
            
            # Employee Data
            if system_field == 'EMPLOYEE_NAME': val = employee.full_name.upper()
            elif system_field == 'EMPLOYEE_DNI': val = employee.dni or ""
            elif system_field == 'EMPLOYEE_EMAIL': val = employee.email or ""
            elif system_field == 'EMPLOYEE_POSITION': val = (employee.position or "").upper()
            elif system_field == 'EMPLOYEE_AREA': val = (employee.area or "").upper()
            elif system_field == 'EMPLOYEE_COMPANY': val = (employee.company or "TRANSTOTAL AGENCIA MARITIMA S.A.").upper()
            elif system_field == 'EMPLOYEE_LOCATION': val = (employee.location or "").upper()
            
            # Date Data
            elif system_field == 'CURRENT_DATE': val = datetime.now().strftime("%d/%m/%Y")
            elif system_field == 'CURRENT_DATE_LONG': val = date_str
            
            # Dynamic Content
            elif system_field == 'ACTA_OBSERVATIONS': val = observations if observations else ""
            
            # Tables
            elif system_field == 'RETURNED_DEVICES_TABLE' or system_field == 'DEVICE_TABLE':
                # Build Table
                headers = ["Nro.", "EQUIPO", "MARCA", "MODELO", "Nro. SERIE"]
                rows = []
                row_num = 1
                for a in assignments:
                    d = a.device
                    row = [
                        str(row_num),
                        d.device_type.upper(),
                        (d.brand or "").upper(),
                        (d.model or "").upper(),
                        (d.serial_number or "").upper()
                    ]
                    rows.append(row)
                    row_num += 1
                    
                    # Implicit Accessories
                    dtype = d.device_type.lower() if d.device_type else ""
                    if dtype == 'monitor':
                        rows.append([str(row_num), "CABLE DE PODER", "-", "-", "-"])
                        row_num += 1
                        rows.append([str(row_num), "CABLE HDMI", "-", "-", "-"])
                        row_num += 1
                
                # Generate XML
                table_map[placeholder] = doc_utils.get_table_xml(headers, rows)
                val = "" # Ensure placeholder is present in final_vars so doc_utils iterates over it
                
            elif system_field == 'MOBILE_DEVICES_TABLE':
                # Build Mobile Devices Table
                # Columns: CANT, EQUIPO, OPERADOR, ESTADO, MARCA, MODELO, NUMERO SERIE, NUMERO IMEI, NUMERO LINEA
                rows = []
                row_num = 1
                for a in assignments:
                    d = a.device
                    dtype = d.device_type.lower() if d.device_type else ""
                    # Only include mobile and chip devices
                    if dtype not in ['mobile', 'chip']:
                        continue
                    
                    row = [
                        str(row_num),
                        "CELULAR" if dtype == "mobile" else "CHIP",
                        (d.carrier or "-").upper(),
                        (d.status or "USADO").upper(),
                        (d.brand or "-").upper(),
                        (d.model or "-").upper(),
                        (d.serial_number or "-").upper(),
                        (d.imei or "-").upper(),
                        (d.phone_number or "-")
                    ]
                    rows.append(row)
                    row_num += 1
                
                # Generate XML with red headers
                table_map[placeholder] = doc_utils.get_mobile_table_xml(rows)
                val = ""
                 
                
            final_vars[placeholder] = val

        # Replace in document
        success = doc_utils.replace_variables_in_docx(template_path, filepath, final_vars, table_map)
        
        if success:
            return filepath
        else:
            print("Error generating dynamic acta, falling back to legacy.")

    # --- FALLBACK TO LEGACY HARDCODED GENERATION ---

def generate_return_acta_mobile(termination_id, employee, assignments, observations=None, template_path=None, template_variables=None):
    # Generate return acta for mobile equipment.
    # If template_path is provided, use dynamic rendering via doc_utils.
    if template_path and os.path.exists(template_path) and template_variables:
        filename = f"Acta_Recepcion_Celular_{termination_id}_{employee.full_name.replace(' ', '_')}.docx"
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "actas")
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, filename)
        
        # Prepare Data Map
        final_vars = {}
        table_map = {}
        date_str = get_spanish_date()
        
        for placeholder, system_field in template_variables.items():
            val = ""
            
            # Employee Data
            if system_field == 'EMPLOYEE_NAME': val = employee.full_name.upper()
            elif system_field == 'EMPLOYEE_DNI': val = employee.dni or ""
            elif system_field == 'EMPLOYEE_EMAIL': val = employee.email or ""
            elif system_field == 'EMPLOYEE_POSITION': val = (employee.position or "").upper()
            elif system_field == 'EMPLOYEE_AREA': val = (employee.area or "").upper()
            elif system_field == 'EMPLOYEE_COMPANY': val = (employee.company or "TRANSTOTAL AGENCIA MARITIMA S.A.").upper()
            elif system_field == 'EMPLOYEE_LOCATION': val = (employee.location or "").upper()
            
            # Date Data
            elif system_field == 'CURRENT_DATE': val = datetime.now().strftime("%d/%m/%Y")
            elif system_field == 'CURRENT_DATE_LONG': val = date_str
            
            # Dynamic Content
            elif system_field == 'ACTA_OBSERVATIONS': val = observations if observations else ""
            
            # Mobile Devices Table
            elif system_field == 'MOBILE_DEVICES_TABLE':
                import json
                rows = []
                row_num = 1
                for a in assignments:
                    d = a.device
                    dtype = d.device_type.lower() if d.device_type else ""
                    
                    # Handle mobile and chip devices
                    if dtype in ["mobile", "chip"]:
                        row = [
                            str(row_num),
                            "CELULAR" if dtype == "mobile" else "CHIP",
                            (d.carrier or "-").upper(),
                            (d.status or "USADO").upper(),
                            (d.brand or "-").upper(),
                            (d.model or "-").upper(),
                            (d.serial_number or "-").upper(),
                            (d.imei or "-").upper(),
                            (d.phone_number or "-")
                        ]
                        rows.append(row)
                        row_num += 1
                        
                        # Add charger row if mobile device has charger details stored in it
                        if dtype == "mobile" and d.mobile_charger_brand:
                            charger_row = [
                                str(row_num),
                                "CARGADOR",
                                "-",
                                "-",
                                (d.mobile_charger_brand or "-").upper(),
                                (d.mobile_charger_model or "-").upper(),
                                (d.mobile_charger_serial or "-").upper(),
                                "-",
                                "-"
                            ]
                            rows.append(charger_row)
                            row_num += 1
                        
                        # Add accessories from specifications if mobile device
                        if dtype == "mobile" and d.specifications:
                            try:
                                specs = json.loads(d.specifications) if isinstance(d.specifications, str) else d.specifications
                                
                                # Define accessories to look for
                                accessories = [
                                    ("Cable USB", "Cable USB"),
                                    ("Auriculares", "Auriculares"),
                                    ("Adaptador Auriculares", "Adaptador Auriculares"),
                                    ("Case", "Case")
                                ]
                                
                                for acc_name, acc_key in accessories:
                                    if acc_key in specs and specs[acc_key] and specs[acc_key].lower() != "no":
                                        acc_value = specs[acc_key]
                                        acc_row = [
                                            str(row_num),
                                            acc_name.upper(),
                                            "-",
                                            "-",
                                            "-",
                                            acc_value.upper() if acc_value else "INCLUIDO",
                                            "-",
                                            "-",
                                            "-"
                                        ]
                                        rows.append(acc_row)
                                        row_num += 1
                            except (json.JSONDecodeError, TypeError):
                                pass  # Skip if specifications can't be parsed
                    
                    # Handle charger devices assigned separately
                    elif dtype == "charger":
                        charger_row = [
                            str(row_num),
                            "CARGADOR",
                            "-",
                            "-",
                            (d.brand or "-").upper(),
                            (d.model or "-").upper(),
                            (d.serial_number or "-").upper(),
                            "-",
                            "-"
                        ]
                        rows.append(charger_row)
                        row_num += 1
                
                table_map[placeholder] = doc_utils.get_mobile_table_xml(rows)
                val = ""
                
            final_vars[placeholder] = val

        # Replace in document
        success = doc_utils.replace_variables_in_docx(template_path, filepath, final_vars, table_map)
        
        if success:
            return filepath
        else:
            print("Error generating dynamic mobile acta, falling back to legacy.")



    # --- FALLBACK TO LEGACY HARDCODED GENERATION ---

    doc = Document()
    
    title = doc.add_paragraph()
    title_run = title.add_run("ACTA DE RECEPCION DE EQUIPOS")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("CESE LABORAL")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph(f"Empleado: {employee.full_name}")
    doc.add_paragraph(f"DNI: {employee.dni or 'N/A'}")
    doc.add_paragraph(f"Correo: {employee.email}")
    doc.add_paragraph(f"Fecha de Cese: {datetime.now().strftime('%d/%m/%Y')}")
    
    doc.add_paragraph()
    
    equip = doc.add_paragraph()
    equip.add_run("EQUIPOS DEVUELTOS:").bold = True
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ['Tipo', 'Marca', 'Modelo', 'Serie', 'Estado']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for assignment in assignments:
        device = assignment.device
        row_cells = table.add_row().cells
        row_cells[0].text = device.device_type.upper()
        row_cells[1].text = device.brand
        row_cells[2].text = device.model
        row_cells[3].text = device.serial_number
        status = device.status
        row_cells[4].text = status.upper() if hasattr(status, 'upper') else str(status).upper()
    
    doc.add_paragraph()
    
    if observations:
        obs = doc.add_paragraph()
        obs.add_run("OBSERVACIONES:").bold = True
        doc.add_paragraph(observations)
        doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.rows[0].cells[0].text = "_" * 30
    sig_table.rows[0].cells[1].text = "_" * 30
    sig_table.rows[1].cells[0].text = "Firma del Empleado"
    sig_table.rows[1].cells[1].text = "Firma del Representante"
    sig_table.rows[2].cells[0].text = employee.full_name
    sig_table.rows[2].cells[1].text = "TRANSTOTAL AGENCIA MARITIMA S.A."
    
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "actas")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"Acta_Recepcion_Computadora_{termination_id}_{employee.full_name.replace(' ', '_')}.docx")
    doc.save(output_path)
    
    return output_path

# Force reload - timestamp: 2026-01-21 12:01
